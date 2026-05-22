import os
import json
import argparse
from docx import Document
from docx.shared import Pt, RGBColor

def parse_color(color_str):
    try:
        parts = [int(x.strip()) for x in color_str.split(",")]
        if len(parts) == 3:
            return RGBColor(*parts)
    except Exception:
        pass
    print("Warning: Invalid color format. Using default dark blue.")
    return RGBColor(31, 78, 121)

def inject_translations(src_docx, trans_json_path, output_docx, para_size, table_size, color_rgb, italic):
    if not os.path.exists(src_docx):
        print(f"Error: Source file not found: {src_docx}")
        return False
    if not os.path.exists(trans_json_path):
        print(f"Error: Translation JSON not found: {trans_json_path}")
        return False

    with open(trans_json_path, 'r', encoding='utf-8') as f:
        translations = json.load(f)

    # 将翻译列表转换成易于检索的字典
    para_trans = {}
    table_trans = {}
    for item in translations:
        if item["type"] == "paragraph":
            para_trans[item["index"]] = item.get("en", "")
        elif item["type"] == "table_cell":
            key = (item["table_idx"], item["row_idx"], item["col_idx"], item["para_idx"])
            table_trans[key] = item.get("en", "")

    doc = Document(src_docx)

    # 1. 回写普通段落
    for idx, para in enumerate(doc.paragraphs):
        if idx in para_trans:
            translation = para_trans[idx].strip()
            if translation:
                run = para.add_run("\n" + translation)
                run.font.size = Pt(para_size)
                run.font.color.rgb = color_rgb
                run.italic = italic

    # 2. 回写表格中的文字，并执行严格的表格级全局合并单元格去重
    for t_idx, table in enumerate(doc.tables):
        processed_cells = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = cell._tc
                if cell_id in processed_cells:
                    continue
                processed_cells.add(cell_id)

                for p_idx, para in enumerate(cell.paragraphs):
                    key = (t_idx, r_idx, c_idx, p_idx)
                    if key in table_trans:
                        translation = table_trans[key].strip()
                        if translation:
                            run = para.add_run("\n" + translation)
                            run.font.size = Pt(table_size)
                            run.font.color.rgb = color_rgb
                            run.italic = italic

    doc.save(output_docx)
    print("Bilingual injection completed successfully!")
    print(f"Bilingual document saved to: {output_docx}")
    return True

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(script_dir)
    
    parser = argparse.ArgumentParser(description="Inject English translations back into a Word (.docx) file.")
    parser.add_argument("-i", "--input", default="/Users/shenweitao/Desktop/印尼服装SRM系统操作手册.docx",
                        help="Path to the input DOCX file")
    parser.add_argument("-t", "--translation", default=os.path.join(parent_dir, "examples", "sample_result.json"),
                        help="Path to the translation JSON file")
    parser.add_argument("-o", "--output", default="/Users/shenweitao/Desktop/印尼服装SRM系统操作手册_双语.docx",
                        help="Path to save the bilingual output DOCX")
    
    # 样式配置参数
    parser.add_argument("--para-size", type=float, default=10.0,
                        help="Font size in pt for normal paragraph translations (default: 10.0)")
    parser.add_argument("--table-size", type=float, default=9.5,
                        help="Font size in pt for table cell translations (default: 9.5)")
    parser.add_argument("--color", default="31,78,121",
                        help="RGB color for English text in format 'R,G,B' (default: '31,78,121' [Dark Blue])")
    parser.add_argument("--no-italic", action="store_true",
                        help="Disable italic styling for English translations")
    
    args = parser.parse_args()
    
    color_rgb = parse_color(args.color)
    italic = not args.no_italic
    
    inject_translations(args.input, args.translation, args.output, args.para_size, args.table_size, color_rgb, italic)
