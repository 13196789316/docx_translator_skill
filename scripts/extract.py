import os
import json
import re
import argparse
from docx import Document

def contains_chinese(text):
    if not text:
        return False
    return bool(re.search(r'[\u4e00-\u9fa5]', text))

def extract_docx_text(docx_path, output_json_path):
    if not os.path.exists(docx_path):
        print(f"Error: File not found at {docx_path}")
        return False

    doc = Document(docx_path)
    extracted_data = []
    item_id = 0

    # 1. 提取普通段落
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if contains_chinese(text):
            extracted_data.append({
                "id": item_id,
                "type": "paragraph",
                "index": idx,
                "text": text
            })
            item_id += 1

    # 2. 提取表格中的文字，并执行表格级全局去重（针对跨行/跨列合并单元格）
    for t_idx, table in enumerate(doc.tables):
        processed_cells = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = cell._tc
                if cell_id in processed_cells:
                    continue
                processed_cells.add(cell_id)
                
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if contains_chinese(text):
                        extracted_data.append({
                            "id": item_id,
                            "type": "table_cell",
                            "table_idx": t_idx,
                            "row_idx": r_idx,
                            "col_idx": c_idx,
                            "para_idx": p_idx,
                            "text": text
                        })
                        item_id += 1

    with open(output_json_path, 'w', encoding='utf-8') as f:
        json.dump(extracted_data, f, ensure_ascii=False, indent=2)

    print(f"Extraction completed. Extracted {len(extracted_data)} items containing Chinese.")
    print(f"Saved source text to {output_json_path}")
    return True

if __name__ == "__main__":
    # 解析命令行参数
    parser = argparse.ArgumentParser(description="Extract Chinese text blocks from a Word (.docx) file.")
    
    # 动态获取桌面路径
    default_input = os.path.join(os.path.expanduser("~"), "Desktop", "manual.docx")
    parser.add_argument("-i", "--input", default=default_input,
                        help="Path to the input DOCX file (default: ~/Desktop/manual.docx)")
    parser.add_argument("-o", "--output", default=None,
                        help="Path to save the extracted source JSON file")
    
    args = parser.parse_args()
    
    # 默认输出路径设定
    if args.output is None:
        # 默认保存在脚本父目录的 examples/sample_source.json 中
        script_dir = os.path.dirname(os.path.abspath(__file__))
        args.output = os.path.join(os.path.dirname(script_dir), "examples", "sample_source.json")
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(args.output), exist_ok=True)
    
    extract_docx_text(args.input, args.output)
